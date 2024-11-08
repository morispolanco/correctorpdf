import streamlit as st
import docx2txt
from io import BytesIO
import requests
import difflib
import html
from docx import Document
from docx.shared import RGBColor

# Configuración de la página
st.set_page_config(
    page_title="Corrector de Textos en Español",
    layout="wide",
)

# Título de la aplicación
st.title("Corrector de Gramática y Estilo para Textos en Español")

# Instrucciones
st.markdown("""
Esta aplicación permite subir un archivo en formato **DOCX** y corregirlo en hasta 200,000 caracteres. 
Selecciona el nivel de corrección y visualiza las correcciones realizadas en el texto con resaltados de colores, similar al "Control de Cambios" de Word.
""")

# Carga de la clave API desde los secretos de Streamlit
try:
    OPENROUTER_API_KEY = st.secrets["openrouter"]["api_key"]
except KeyError:
    st.error("Clave API de OpenRouter no encontrada. Por favor, configúrala en los secretos de Streamlit.")
    st.stop()

# Función para extraer texto del DOCX
def extract_text_from_docx(file):
    try:
        # Guardar el archivo en memoria
        with BytesIO() as docx_buffer:
            docx_buffer.write(file.read())
            docx_buffer.seek(0)
            text = docx2txt.process(docx_buffer)
        return text
    except Exception as e:
        st.error(f"Error al extraer texto del DOCX: {e}")
        return ""

# Función para dividir el texto en chunks si es necesario
def split_text(text, max_length=1500):
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = ""
    for para in paragraphs:
        if len(current_chunk) + len(para) + 1 <= max_length:
            current_chunk += para + "\n"
        else:
            chunks.append(current_chunk)
            current_chunk = para + "\n"
    if current_chunk:
        chunks.append(current_chunk)
    return chunks

# Función para corregir el texto usando la API de OpenRouter
def correct_text(text, level):
    if level == "Mínimo":
        prompt = "Corrige la ortografía y la puntuación del siguiente texto en español:\n\n" + text
    elif level == "Intermedio":
        prompt = "Corrige la ortografía, gramática y puntuación del siguiente texto en español:\n\n" + text
    elif level == "Avanzado":
        prompt = "Corrige la ortografía, gramática y puntuación del siguiente texto en español. Además, mejora el estilo y ofrece sugerencias de redacción:\n\n" + text
    else:
        prompt = text  # Default to no correction

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENROUTER_API_KEY}"
    }

    data = {
        "model": "openai/gpt-4o-mini",
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ]
    }

    try:
        response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=data)
        response.raise_for_status()
        result = response.json()
        return result['choices'][0]['message']['content']
    except requests.exceptions.RequestException as e:
        st.error(f"Error en la API: {e}")
        return None
    except KeyError:
        st.error("Respuesta inesperada de la API.")
        return None

# Función para resaltar las diferencias entre el texto original y el corregido (a nivel de palabra)
def highlight_differences_word_level(original, corrected):
    # Escapar caracteres HTML
    original = html.escape(original)
    corrected = html.escape(corrected)

    # Dividir en palabras
    original_words = original.split()
    corrected_words = corrected.split()

    # Crear una instancia de Differ
    differ = difflib.Differ()
    diff = list(differ.compare(original_words, corrected_words))

    # Construir el HTML con resaltados
    html_diff = ""
    for word in diff:
        if word.startswith("- "):
            # Eliminaciones en rojo con tachado
            html_diff += f'<span style="background-color: #ffcccc; text-decoration: line-through;">{word[2:]}</span> '
        elif word.startswith("+ "):
            # Inserciones en verde
            html_diff += f'<span style="background-color: #ccffcc;">{word[2:]}</span> '
        else:
            # Palabras sin cambios
            html_diff += f'{word[2:]} '
    return html_diff

# Función para crear un DOCX con control de cambios simulado
def create_docx_with_changes_word_level(original, corrected):
    document = Document()

    # Usar difflib para comparar palabras
    differ = difflib.Differ()
    diff = list(differ.compare(original.split(), corrected.split()))

    para = document.add_paragraph()
    for word in diff:
        if word.startswith("- "):
            # Eliminación: palabra eliminada
            run = para.add_run(word[2:] + " ")
            run.font.color.rgb = RGBColor(255, 0, 0)  # Rojo
            run.font.strike = True  # Tachado
        elif word.startswith("+ "):
            # Inserción: palabra añadida
            run = para.add_run(word[2:] + " ")
            run.font.color.rgb = RGBColor(0, 128, 0)  # Verde
        else:
            # Palabra sin cambios
            para.add_run(word[2:] + " ")

    # Guardar el documento en un objeto BytesIO
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

# Interfaz de usuario para subir el archivo
uploaded_file = st.file_uploader("Sube tu archivo DOCX", type=["docx"])

# Selección del nivel de corrección
correction_level = st.selectbox(
    "Selecciona el nivel de corrección",
    ("Mínimo (Ortografía y Puntuación)", "Intermedio (Ortografía, Gramática y Puntuación)", "Avanzado (Incluye Estilo y Redacción)")
)

if uploaded_file is not None:
    with st.spinner("Extrayendo texto del archivo..."):
        text = extract_text_from_docx(uploaded_file)

    if not text:
        st.error("No se pudo extraer texto del archivo.")
    elif len(text) > 200000:
        st.error("El documento supera los 200,000 caracteres. Por favor, sube un archivo más pequeño.")
    else:
        st.success("Texto extraído exitosamente.")
        st.subheader("Texto Original")
        st.text_area("Texto Original", text, height=300)
        
        if st.button("Corregir Texto"):
            with st.spinner("Corrigiendo el texto..."):
                # Determinar el nivel de corrección
                if correction_level.startswith("Mínimo"):
                    level = "Mínimo"
                elif correction_level.startswith("Intermedio"):
                    level = "Intermedio"
                else:
                    level = "Avanzado"

                # Dividir el texto en chunks si es muy largo
                chunks = split_text(text)
                corrected_chunks = []

                # Inicializar la barra de progreso
                progress_bar = st.progress(0)
                total_chunks = len(chunks)

                for i, chunk in enumerate(chunks):
                    # Actualizar la barra de progreso
                    progress = (i + 1) / total_chunks
                    progress_bar.progress(progress)

                    st.write(f"Corrigiendo fragmento {i+1} de {total_chunks}...")
                    corrected = correct_text(chunk, level)
                    if corrected:
                        corrected_chunks.append(corrected)
                    else:
                        st.error("Error al corregir el texto.")
                        break

                corrected_text = "\n".join(corrected_chunks)

                # Completar la barra de progreso
                progress_bar.progress(1.0)

                if corrected_text:
                    st.success("Texto corregido exitosamente.")
                    st.subheader("Texto Corregido")
                    st.text_area("Texto Corregido", corrected_text, height=300)

                    st.subheader("Control de Cambios")
                    st.markdown(
                        highlight_differences_word_level(text, corrected_text),
                        unsafe_allow_html=True
                    )

                    # Crear DOCX con control de cambios simulado
                    docx_buffer = create_docx_with_changes_word_level(text, corrected_text)

                    # Opción para descargar el DOCX corregido
                    st.download_button(
                        label="Descargar Documento Corregido (DOCX)",
                        data=docx_buffer,
                        file_name="corregido.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    # Opción para descargar el texto corregido en TXT
                    corrected_text_bytes = corrected_text.encode('utf-8')
                    st.download_button(
                        label="Descargar Texto Corregido (TXT)",
                        data=corrected_text_bytes,
                        file_name="corregido.txt",
                        mime="text/plain"
                    )
